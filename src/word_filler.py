from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document


ANCHORS = {
    "20 tháng 04 năm 2026": "document_date_long",
    "0000001/GCN/001-P.KD12/HH.1.1.4/2026": "policy_numbers",
    "Hóa chất Methyl ethyl ketone": "commodities",
    "Methyl ethyl ketone": "commodities",
    "7,068,600 VND": "survey_fee",
    "26NK0742/BH": "survey_report_number",
    "JINZHOU STAR": "vessel_name",
    "JSV2602CS01MEK": "bill_of_lading_numbers",
    "31/01/2026": "arrival_date",
    "790.297": "bl_quantity_mt",
    "788.290": "received_quantity_mt",
    "- 2.007": "shortage_mt_spaced",
    "- 0.254": "shortage_percent_spaced",
    "T107": "shore_tanks",
    "số 119": "invoice_number_phrase",
    "12/02/2026": "invoice_date",
    "20/04/2026": "document_date",
    "Trung Quốc": "origin",
}

ALIASES = {
    "document_date_long": "document_date",
    "invoice_number_phrase": "invoice_number",
    "shortage_mt_spaced": "shortage_mt",
    "shortage_percent_spaced": "shortage_percent",
}


def editable_field_names() -> set[str]:
    return {ALIASES.get(name, name) for name in ANCHORS.values()}


def layout_risks(values: dict[str, str], ratio: float = 2.2) -> list[str]:
    """Warn when replacement text is much longer than the highlighted sample."""
    display = _display(values)
    risks: list[str] = []
    checked: set[str] = set()
    for old, output_name in ANCHORS.items():
        canonical = ALIASES.get(output_name, output_name)
        if canonical in checked:
            continue
        checked.add(canonical)
        new = display.get(output_name, "")
        if new and len(new) > max(len(old) * ratio, len(old) + 25):
            risks.append(canonical)
    return risks


def _is_yellow(run) -> bool:
    from docx.oxml.ns import qn

    rpr = run._element.rPr
    if rpr is None:
        return False
    highlight = rpr.find(qn("w:highlight"))
    return highlight is not None and highlight.get(qn("w:val")) == "yellow"


def _replace_across_runs(paragraph, old: str, new: str, yellow_only: bool = True) -> int:
    if old not in paragraph.text:
        return 0
    full = paragraph.text
    count = full.count(old)
    start = 0
    for _ in range(count):
        start = full.find(old, start)
        end = start + len(old)
        positions: list[tuple[int, int, int]] = []
        cursor = 0
        for index, run in enumerate(paragraph.runs):
            run_end = cursor + len(run.text)
            if run_end > start and cursor < end:
                positions.append((index, max(start - cursor, 0), min(end - cursor, len(run.text))))
            cursor = run_end
        if not positions:
            break
        if yellow_only and not all(_is_yellow(paragraph.runs[idx]) for idx, _, _ in positions):
            start = end
            continue
        first_idx, first_start, first_end = positions[0]
        first = paragraph.runs[first_idx]
        suffix = paragraph.runs[positions[-1][0]].text[positions[-1][2]:]
        first.text = first.text[:first_start] + new + suffix
        for idx, _, _ in positions[1:]:
            paragraph.runs[idx].text = ""
        full = paragraph.text
        start += len(new)
    return count


def _iter_paragraphs(document):
    seen = set()
    containers = [document]
    for section in document.sections:
        containers.extend([section.header, section.footer, section.first_page_header, section.first_page_footer])
    for container in containers:
        for paragraph in container.paragraphs:
            if paragraph._p not in seen:
                seen.add(paragraph._p); yield paragraph
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph._p not in seen:
                            seen.add(paragraph._p); yield paragraph


def _display(values: dict[str, str]) -> dict[str, str]:
    out = {key: str(value or "") for key, value in values.items()}
    date = out.get("document_date", "")
    parts = date.split("/")
    out["document_date_long"] = f"{parts[0]} tháng {parts[1]} năm {parts[2]}" if len(parts) == 3 else date
    out["policy_numbers"] = "; ".join(filter(None, re.split(r"\s*[;\n]+\s*", out.get("policy_numbers", ""))))
    out["commodities"] = "; ".join(filter(None, re.split(r"\s*[;\n]+\s*", out.get("commodities", ""))))
    out["bill_of_lading_numbers"] = "; ".join(filter(None, re.split(r"\s*[;\n]+\s*", out.get("bill_of_lading_numbers", ""))))
    out["shore_tanks"] = ", ".join(filter(None, re.split(r"\s*[;\n]+\s*", out.get("shore_tanks", ""))))
    out["discharge_port_phrase"] = out.get("discharge_port", "")
    out["invoice_number_phrase"] = f"số {out.get('invoice_number', '')}".rstrip()
    out["shortage_mt_spaced"] = out.get("shortage_mt", "").replace("-", "- ")
    out["shortage_percent_spaced"] = out.get("shortage_percent", "").replace("-", "- ")
    return out


def fill_template(template_path: str | Path | bytes, values: dict[str, str]) -> bytes:
    document = Document(BytesIO(template_path) if isinstance(template_path, bytes) else str(template_path))
    display = _display(values)
    for paragraph in _iter_paragraphs(document):
        for old, field_name in sorted(ANCHORS.items(), key=lambda item: len(item[0]), reverse=True):
            if field_name in display:
                _replace_across_runs(paragraph, old, display[field_name], yellow_only=True)
        for name, value in display.items():
            _replace_across_runs(paragraph, "{{" + name + "}}", value, yellow_only=True)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def remaining_placeholders(data: bytes) -> list[str]:
    document = Document(BytesIO(data))
    text = "\n".join(p.text for p in _iter_paragraphs(document))
    return sorted(set(re.findall(r"{{([a-zA-Z0-9_]+)}}", text)))


def remove_all_highlights(data: bytes) -> bytes:
    """Remove Word text highlight without changing any other formatting."""
    from docx.oxml.ns import qn

    document = Document(BytesIO(data))
    for paragraph in _iter_paragraphs(document):
        for run in paragraph.runs:
            rpr = run._element.rPr
            if rpr is None:
                continue
            for highlight in list(rpr.findall(qn("w:highlight"))):
                rpr.remove(highlight)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
